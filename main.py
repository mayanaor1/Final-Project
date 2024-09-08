import tkinter as tk
from tkinter import messagebox, scrolledtext, filedialog
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import undetected_chromedriver as uc
from fake_useragent import UserAgent
from time import sleep
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from docx import Document
import os

def OpenGPT(username, password):
    """
    Initialize and open a Chrome WebDriver session for ChatGPT.

    This function sets up the Chrome WebDriver,
    navigates to the ChatGPT website, and logs in using the provided credentials.

    Args:
    username (str): The username for ChatGPT login.
    password (str): The password for ChatGPT login.

    Returns:
    WebDriver or None: The initialized WebDriver if successful, None otherwise.
    """
    # Initialize Chrome WebDriver options
    op = webdriver.ChromeOptions()
    op.add_argument(f"user-agent={UserAgent.random}")  # Use a random user agent
    op.add_argument("user-data-dir=./")  # Set user data directory
    op.add_experimental_option("detach", True)  # Keep browser open after script finishes
    op.add_experimental_option("excludeSwitches", ["enable-logging"])  # Disable logging

    try:
        driver = uc.Chrome(chrome_options=op)
    except Exception as e:
        print("Error initializing WebDriver:", e)
        return None

    try:
        driver.get('https://chatgpt.com/')  # Navigate to ChatGPT website
        driver.maximize_window()

        # Wait for and click the "Continue" button
        continue_button = WebDriverWait(driver, 50).until(
            EC.element_to_be_clickable((By.CLASS_NAME, "btn-primary"))
        )
        continue_button.click()

        # Enter username
        username_field = WebDriverWait(driver, 50).until(
            EC.visibility_of_element_located((By.ID, 'email-input'))
        )
        username_field.send_keys(username)

        continue_button = driver.find_element(By.CLASS_NAME, "continue-btn")
        continue_button.click()

        # Enter password
        password_field = WebDriverWait(driver, 50).until(
            EC.visibility_of_element_located((By.ID, 'password'))
        )
        password_field.send_keys(password)
        submit_button = driver.find_element(By.CLASS_NAME, "_button-login-password")
        submit_button.click()

        sleep(50)  # Wait for page to load completely

        return driver
    except (TimeoutException, NoSuchElementException) as e:
        print("Error during login:", e)
        driver.quit()
        return None

def ChatGPT(question):
    """
    Send a question to ChatGPT and retrieve the response.

    This function interacts with the ChatGPT interface, sends the provided question,
    and waits for a response.

    Args:
    question (str): The question to be sent to ChatGPT.

    Returns:
    str or None: The response from ChatGPT if successful, None otherwise.
    """
    try:
        # Start a new chat
        sleep(2)
        new_button = WebDriverWait(driver, 50).until(
            EC.element_to_be_clickable((By.CLASS_NAME, 'text-token-text-primary'))
        )
        new_button.click()

        sleep(2)  # Wait for new chat interface to load

        # Modify the question
        modified_question = f"{question} answer in one sentence and end your answer with the word bye"

        # Enter and send the modified question
        chat_input = WebDriverWait(driver, 50).until(
            EC.visibility_of_element_located((By.ID, 'prompt-textarea'))
        )
        chat_input.send_keys(modified_question)
        print("Sent question:", modified_question)
        chat_input.send_keys(Keys.RETURN)

        # Wait for response
        WebDriverWait(driver, 30).until(
            lambda driver: 'bye' in driver.find_element(By.CLASS_NAME, 'agent-turn').text.lower()
        )

        print("Waiting for response")
        response = WebDriverWait(driver, 30).until(
            EC.visibility_of_element_located((By.CLASS_NAME, 'agent-turn'))
        )
        print("Received response")

        # Remove "bye" and everything after it from the response
        response_text = response.text
        index = response_text.lower().rfind('bye')
        if index != -1:
            response_text = response_text[:index]

        return response_text.strip()
    except (TimeoutException, NoSuchElementException) as e:
        error_message = driver.find_element(By.CLASS_NAME, "text-token-text-error")
        if error_message:
            print("You've reached our limit of messages per hour. Please try again later.")
            driver.quit()
            exit(1)
        else:
            print("Error in ChatGPT interaction:", e)
            driver.quit()
        return None

def SaveDocx(questions, answers, save_path):
    """
    Convert the questions and answers to a Word document.

    This function creates a formatted Word document with questions and answers.

    Args:
    questions (list): List of questions.
    answers (list): List of lists containing answers for each question.
    save_path (str): Path to save the document.

    Returns:
    str or None: The path to the new Word document if successful, None otherwise.
    """
    try:
        # Create a new Word document
        doc = Document()

        for question, answer_list in zip(questions, answers):
            # Add formatted question
            title1 = doc.add_paragraph("Question:")
            title1.runs[0].bold = True
            title1.runs[0].underline = True

            doc.add_paragraph(question)

            # Add formatted answers
            title2 = doc.add_paragraph("Answers:")
            title2.runs[0].bold = True
            title2.runs[0].underline = True

            for answer in answer_list:
                doc.add_paragraph(str(answer))

            doc.add_page_break()

        # Save the Word document
        docx_path = os.path.join(save_path, 'ChatGPT_Answers.docx')
        doc.save(docx_path)
        print("Saved document successfully")

        return docx_path

    except Exception as e:
        print("Error saving document:", e)
        return None

def choose_save_location():
    """
    Open a dialog for the user to choose a save location.

    Returns:
    str: The selected directory path.
    """
    directory = filedialog.askdirectory()
    if directory:
        save_location_entry.delete(0, tk.END)
        save_location_entry.insert(0, directory)

def run_script():
    """
    Main function to run the ChatGPT automation script.

    This function is triggered when the user clicks the "Start" button in the GUI.
    It orchestrates the entire process of interacting with ChatGPT, saving responses,
    and creating the final Word document.
    """
    username = username_entry.get()
    password = password_entry.get()
    iterations = int(iterations_spinbox.get())
    questions = questions_text.get("1.0", tk.END).strip().split('\n')
    save_location = save_location_entry.get()

    if not username or not password:
        messagebox.showerror("Input Error", "Please enter both username and password.")
        return

    if not questions or questions == [""]:
        messagebox.showerror("Input Error", "Please enter at least one question.")
        return

    if not save_location:
        messagebox.showerror("Input Error", "Please select a save location.")
        return

    global driver
    driver = OpenGPT(username, password)

    if driver:
        try:
            all_answers = [[] for _ in questions]

            for i in range(iterations):
                print(f"Starting iteration {i + 1}/{iterations}...")

                for idx, question in enumerate(questions):
                    answer = ChatGPT(question)
                    all_answers[idx].append(answer if answer else "No response")

                print(f"Iteration {i + 1} completed.")

            docx_path = SaveDocx(questions, all_answers, save_location)

            driver.quit()

            if docx_path:
                messagebox.showinfo("Success", f"Answers saved to Word document: {docx_path}")
            else:
                messagebox.showerror("Error", "Failed to save answers to Word document.")

        except Exception as e:
            driver.quit()
            messagebox.showerror("Error", f"Error processing questions: {e}")

# Create the main window
root = tk.Tk()
root.title("ChatGPT Automation")

# Create and place widgets
tk.Label(root, text="ChatGPT username:").grid(row=0, column=0, padx=10, pady=10)
username_entry = tk.Entry(root, width=30)
username_entry.grid(row=0, column=1, padx=10, pady=10)

tk.Label(root, text="ChatGPT password:").grid(row=1, column=0, padx=10, pady=10)
password_entry = tk.Entry(root, width=30, show='*')
password_entry.grid(row=1, column=1, padx=10, pady=10)

tk.Label(root, text="Iterations:").grid(row=2, column=0, padx=10, pady=10)
iterations_spinbox = tk.Spinbox(root, from_=1, to=100, width=5)
iterations_spinbox.grid(row=2, column=1, padx=10, pady=10)

tk.Label(root, text="Questions (one per line):").grid(row=3, column=0, columnspan=2, padx=10, pady=5)
questions_text = scrolledtext.ScrolledText(root, width=50, height=10)
questions_text.grid(row=4, column=0, columnspan=2, padx=10, pady=5)

tk.Label(root, text="Save Location:").grid(row=5, column=0, padx=10, pady=10)
save_location_entry = tk.Entry(root, width=30)
save_location_entry.grid(row=5, column=1, padx=10, pady=10)
choose_location_button = tk.Button(root, text="Choose", command=choose_save_location)
choose_location_button.grid(row=5, column=2, padx=10, pady=10)

start_button = tk.Button(root, text="Start", command=run_script)
start_button.grid(row=6, column=0, columnspan=3, padx=10, pady=10)

# Run the GUI event loop
root.mainloop()